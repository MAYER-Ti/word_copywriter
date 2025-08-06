import re
from docx import Document
import pdfplumber
from pdf2image import convert_from_path
import pytesseract

DEFAULT_DATA = {
    "Данные заказчика": "",
    "ИНН получателя": "",
    "ОГРН получателя": "",
    "Номер документа": "",
    "Адрес загрузки": "",
    "Адрес разгрузки": "",
    "Марка автомобиля": "",
    "Номер полуприцепа": "",
    "ФИО водителя": "",
    "Дата погрузки": "",
    "Дата разгрузки": "",
    "Стоимость перевозки": "",
    "Цена": "",
}


def extract_price(cost: str) -> str:
    """Extract numeric price from a cost string."""
    match = re.search(r"\d[\d\s]*", cost or "")
    return match.group(0).replace(" ", "") if match else ""


def read_data_from_docx(path):
    """Parse contract-like document and return values for placeholders."""
    doc = Document(path)

    data = DEFAULT_DATA.copy()

    # First paragraph with number
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Договор-заявка"):
            num_start = text.find("№")
            if num_start != -1:
                data["Номер документа"] = text[num_start:].strip()
            break

    # First table with route and cost info
    if doc.tables:
        t = doc.tables[0]
        try:
            data["Адрес загрузки"] = t.cell(8, 0).text.strip()
            data["Адрес разгрузки"] = t.cell(8, 4).text.strip()
            data["Дата погрузки"] = t.cell(10, 0).text.strip()
            data["Дата разгрузки"] = t.cell(10, 4).text.strip()
            data["Стоимость перевозки"] = t.cell(11, 4).text.strip()
        except IndexError:
            pass

    # Second table with vehicle info
    if len(doc.tables) > 1:
        t = doc.tables[1]
        try:
            data["Марка автомобиля"] = t.cell(0, 1).text.strip()
            data["Номер полуприцепа"] = t.cell(0, 2).text.strip()
            data["ФИО водителя"] = t.cell(1, 1).text.strip()
        except IndexError:
            pass

    # Third table with customer info
    if len(doc.tables) > 2:
        t = doc.tables[2]
        try:
            cell_text = t.cell(0, 1).text
        except IndexError:
            cell_text = ""
        if cell_text:
            start = cell_text.find("Заказчик:")
            end = cell_text.find("Почтовый адрес")
            if start != -1 and end != -1 and end > start:
                data["Данные заказчика"] = cell_text[start + len("Заказчик:"):end].strip()
            inn_match = re.search(r"ИНН получателя \d+", cell_text)
            if inn_match:
                data["ИНН получателя"] = inn_match.group(0).strip()
            ogrn_match = re.search(r"ОГРН \d+", cell_text)
            if ogrn_match:
                data["ОГРН получателя"] = ogrn_match.group(0).strip()

    data["Цена"] = extract_price(data["Стоимость перевозки"])
    return data


def parse_data_from_text(text: str):
    """Parse plain text of contract-like document and return values."""
    data = DEFAULT_DATA.copy()

    # Split text into non-empty lines for easier table-like parsing
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    def split_cols(line: str):
        """Split a line into columns by tabs or 2+ spaces."""
        return [p.strip() for p in re.split(r"\t+|\s{2,}", line) if p.strip()]

    def find(pattern, flags=0):
        match = re.search(pattern, text, flags)
        return match.group(1).strip() if match else ""

    # Document number can appear anywhere
    num_match = re.search(r"Договор-заявка.*?(№\s*[^\n\r]+)", text)
    if num_match:
        data["Номер документа"] = num_match.group(1).strip()

    for i, line in enumerate(lines):
        cols = split_cols(line)
        if re.match(r"Адрес загрузки\s+Адрес разгрузки", line) and i + 1 < len(lines):
            addr_line = lines[i + 1]
            j = i + 2
            while j < len(lines) and not lines[j].startswith("Дата"):
                addr_line += " " + lines[j]
                j += 1
            split_match = re.search(r"\d\s+(?=[А-ЯЁ])", addr_line)
            if split_match:
                idx = split_match.end()
                data["Адрес загрузки"] = addr_line[:idx].strip()
                data["Адрес разгрузки"] = addr_line[idx:].strip()
            else:
                vals = split_cols(addr_line)
                if len(vals) >= 2:
                    data["Адрес загрузки"], data["Адрес разгрузки"] = vals[0], vals[1]
        elif re.match(r"Дата\s+Время\s+Дата\s+Время", line) and i + 1 < len(lines):
            vals_line = lines[i + 1]
            if i + 2 < len(lines) and not lines[i + 2].startswith("Стоимость перевозки"):
                vals_line += " " + lines[i + 2]
            vals = split_cols(vals_line)
            if len(vals) >= 3:
                data["Дата погрузки"] = vals[0]
                data["Дата разгрузки"] = vals[2]
        elif line.startswith("Стоимость перевозки"):
            cost_match = re.search(r"Стоимость перевозки(?:\s*\(прописью\))?\s*([^\n\r]+)", line)
            if cost_match:
                data["Стоимость перевозки"] = cost_match.group(1).strip()
        elif line.startswith("Марка") and "полуприцеп" in line:
            vals = split_cols(line)
            if len(vals) > 1:
                data["Марка автомобиля"] = vals[1]
            if len(vals) > 2:
                data["Номер полуприцепа"] = vals[2]
        elif line.startswith("ФИО водителя"):
            vals = split_cols(line)
            if len(vals) > 1:
                data["ФИО водителя"] = vals[1]

    # Fallback regex-based extraction if table parsing failed
    if not (data["Адрес загрузки"] and data["Адрес разгрузки"]):
        addr_match = re.search(r"Адрес загрузки[:\s]+([^\n\r]+)\s+Адрес разгрузки[:\s]+([^\n\r]+)", text)
        if addr_match:
            data["Адрес загрузки"], data["Адрес разгрузки"] = addr_match.group(1).strip(), addr_match.group(2).strip()
    if not data["Адрес загрузки"]:
        data["Адрес загрузки"] = find(r"(?m)^Адрес загрузки[:\s]+([^\n\r]+)")
    if not data["Адрес разгрузки"]:
        data["Адрес разгрузки"] = find(r"(?m)^Адрес разгрузки[:\s]+([^\n\r]+)")
    if not (data["Дата погрузки"] and data["Дата разгрузки"]):
        dates = re.findall(r"\d{2}\.\d{2}\.\d{4}", text)
        if len(dates) >= 2:
            data["Дата погрузки"], data["Дата разгрузки"] = dates[0], dates[1]
    if not data["Стоимость перевозки"]:
        cost_match = re.search(r"Стоимость перевозки(?:\s*\(прописью\))?\s*([^\n\r]+)", text)
        if cost_match:
            data["Стоимость перевозки"] = cost_match.group(1).strip()
    if not (data["Марка автомобиля"] and data["Номер полуприцепа"]):
        car_match = re.search(r"Марка[,\s]+номер а/м, номер полуприцепа\s+([^\s\n]+)\s+([^\s\n]+)", text, re.I)
        if car_match:
            data["Марка автомобиля"], data["Номер полуприцепа"] = car_match.group(1).strip(), car_match.group(2).strip()
    if not data["Марка автомобиля"]:
        data["Марка автомобиля"] = find(r"(?m)^Марка автомобиля[:\s]+([^\n\r]+)")
    if not data["Номер полуприцепа"]:
        data["Номер полуприцепа"] = find(r"(?m)^Номер полуприцепа[:\s]+([^\n\r]+)")
    if not data["ФИО водителя"]:
        data["ФИО водителя"] = find(r"(?m)^ФИО водителя[:\s]*([^\n\r]+)")

    cust_matches = re.findall(r"Заказчик:(.*?)(?:Почтовый адрес|$)", text, re.S)
    if not cust_matches:
        cust_matches = re.findall(r"Заказчик\s+(.*?)(?:Почтовый адрес|$)", text, re.S)
    if cust_matches:
        block = cust_matches[-1]
        lines_cust = [ln.strip() for ln in block.splitlines() if ln.strip()]
        name = ""
        address = ""
        idxs = [i for i, ln in enumerate(lines_cust) if ln.startswith("Индивидуальный предприниматель")]
        if idxs:
            idx = idxs[-1]
            after = lines_cust[idx][len("Индивидуальный предприниматель"):].strip()
            if not after or after == "Индивидуальный предприниматель":
                if idx + 1 < len(lines_cust):
                    parts = lines_cust[idx + 1].split()
                    if len(parts) >= 3:
                        name = "Индивидуальный предприниматель " + " ".join(parts[-3:])
            else:
                for delim in ["Юридический адрес", "Почтовый адрес"]:
                    pos = after.find(delim)
                    if pos != -1:
                        after = after[:pos].strip()
                        break
                name = "Индивидуальный предприниматель " + after
        for j, ln in enumerate(lines_cust):
            if ln.startswith("Юридический адрес"):
                address = ln
                for extra in lines_cust[j + 1:]:
                    if "литера" in extra or "офис" in extra:
                        address = address.rstrip(",") + ", " + extra
                        break
                break
        if not address:
            addr_match = re.search(r"Юридический адрес[^\n]*(?=\n|Почтовый адрес|$)", block)
            if addr_match:
                address = addr_match.group(0).strip()
        if name:
            data["Данные заказчика"] = name
            if address:
                data["Данные заказчика"] += "\n" + address

    inn_match = re.search(r"(ИНН получателя\s*\d+)", text)
    if inn_match:
        data["ИНН получателя"] = inn_match.group(1).strip()

    ogrn_match = re.search(r"(ОГРН\s*\d+)", text)
    if ogrn_match:
        data["ОГРН получателя"] = ogrn_match.group(1).strip()

    data["Цена"] = extract_price(data["Стоимость перевозки"])
    return data


def extract_text_from_pdf(path: str) -> str:
    """Extract text from PDF, using OCR for scanned documents."""
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                text += page_text + "\n"
    if text.strip():
        return text

    # Fallback to OCR for scanned PDFs
    images = convert_from_path(path)
    for img in images:
        text += pytesseract.image_to_string(img, lang="rus") + "\n"
    return text


def read_data_from_pdf(path: str):
    text = extract_text_from_pdf(path)
    return parse_data_from_text(text)


def read_data_from_file(path: str):
    if path.lower().endswith(".docx"):
        return read_data_from_docx(path)
    if path.lower().endswith(".pdf"):
        return read_data_from_pdf(path)
    raise ValueError("Unsupported file format")
