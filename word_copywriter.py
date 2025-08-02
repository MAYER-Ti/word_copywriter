import sys
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QFileDialog, QMessageBox
from docx import Document
import re
import pdfplumber
from pdf2image import convert_from_path
import pytesseract


DEFAULT_DATA = {
    "Данные заказчика": "",
    "ИНН получателя": "",
    "ОГРН получателя": "",
    "Номер документа": "",
    "Адрес загрузки": "",
    "Адрес разгрузки": "",
    "Марка автомобиля": "",
    "Номер полуприцепа": "",
    "ФИО водителя": "",
    "Дата погрузки": "",
    "Дата разгрузки": "",
    "Стоимость перевозки": "",
}


def read_data_from_docx(path):
    """Parse contract-like document and return values for placeholders."""
    doc = Document(path)

    data = DEFAULT_DATA.copy()

    # First paragraph with number
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Договор-заявка"):# and text.endswith("года ."):
            num_start = text.find("№")
            if num_start != -1:
                data["Номер документа"] = text[num_start:].strip()
            break

    # First table with route and cost info
    if doc.tables:
        t = doc.tables[0]
        try:
            data["Адрес загрузки"] = t.cell(8, 0).text.strip()
            data["Адрес разгрузки"] = t.cell(8, 4).text.strip()
            data["Дата погрузки"] = t.cell(10, 0).text.strip()
            data["Дата разгрузки"] = t.cell(10, 4).text.strip()
            data["Стоимость перевозки"] = t.cell(11, 4).text.strip()
        except IndexError:
            pass

    # Second table with vehicle info
    if len(doc.tables) > 1:
        t = doc.tables[1]
        try:
            data["Марка автомобиля"] = t.cell(0, 1).text.strip()
            data["Номер полуприцепа"] = t.cell(0, 2).text.strip()
            data["ФИО водителя"] = t.cell(1, 1).text.strip()
        except IndexError:
            pass

    # Third table with customer info
    if len(doc.tables) > 2:
        t = doc.tables[2]
        try:
            cell_text = t.cell(0, 1).text
        except IndexError:
            cell_text = ""
        if cell_text:
            start = cell_text.find("Заказчик:")
            end = cell_text.find("Почтовый адрес")
            if start != -1 and end != -1 and end > start:
                data["Данные заказчика"] = cell_text[start + len("Заказчик:"):end].strip()
            inn_match = re.search(r"ИНН получателя \d+", cell_text)
            if inn_match:
                data["ИНН получателя"] = inn_match.group(0).strip()
            ogrn_match = re.search(r"ОГРН \d+", cell_text)
            if ogrn_match:
                data["ОГРН получателя"] = ogrn_match.group(0).strip()

    return data



def parse_data_from_text(text: str):
    """Parse plain text of contract-like document and return values."""
    data = DEFAULT_DATA.copy()

    # Split text into non-empty lines for easier table-like parsing
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    def split_cols(line: str):
        """Split a line into columns by tabs or 2+ spaces."""
        return [p.strip() for p in re.split(r"\t+|\s{2,}", line) if p.strip()]

    def find(pattern, flags=0):
        match = re.search(pattern, text, flags)
        return match.group(1).strip() if match else ""

    # Document number can appear anywhere
    num_match = re.search(r"Договор-заявка.*?(№\s*[^\n\r]+)", text)
    if num_match:
        data["Номер документа"] = num_match.group(1).strip()

    for i, line in enumerate(lines):
        cols = split_cols(line)
        if cols[:2] == ["Адрес загрузки", "Адрес разгрузки"] and i + 1 < len(lines):
            vals = split_cols(lines[i + 1])
            if len(vals) >= 2:
                data["Адрес загрузки"] = vals[0]
                data["Адрес разгрузки"] = vals[1]
        elif cols[:4] == ["Дата", "Время", "Дата", "Время"] and i + 1 < len(lines):
            vals = split_cols(lines[i + 1])
            if len(vals) >= 3:
                data["Дата погрузки"] = vals[0]
                data["Дата разгрузки"] = vals[2]
        elif line.startswith("Стоимость перевозки"):
            vals = split_cols(line)
            if len(vals) > 1:
                data["Стоимость перевозки"] = vals[1]
        elif line.startswith("Марка") and "полуприцеп" in line:
            vals = split_cols(line)
            if len(vals) > 1:
                data["Марка автомобиля"] = vals[1]
            if len(vals) > 2:
                data["Номер полуприцепа"] = vals[2]
        elif line.startswith("ФИО водителя") or line.startswith("Водитель"):
            vals = split_cols(line)
            if len(vals) > 1:
                data["ФИО водителя"] = vals[1]

    # Fallback regex-based extraction if table parsing failed
    if not data["Адрес загрузки"]:
        data["Адрес загрузки"] = find(r"Адрес загрузки[:\s]*([^\n\r]+)")
    if not data["Адрес разгрузки"]:
        data["Адрес разгрузки"] = find(r"Адрес разгрузки[:\s]*([^\n\r]+)")
    if not data["Дата погрузки"]:
        data["Дата погрузки"] = find(r"Дата погрузки[:\s]*([^\n\r]+)")
    if not data["Дата разгрузки"]:
        data["Дата разгрузки"] = find(r"Дата разгрузки[:\s]*([^\n\r]+)")
    if not data["Стоимость перевозки"]:
        data["Стоимость перевозки"] = find(r"Стоимость перевозки[:\s]*([^\n\r]+)")
    if not data["Марка автомобиля"]:
        data["Марка автомобиля"] = find(r"Марка автомобиля[:\s]*([^\n\r]+)")
    if not data["Номер полуприцепа"]:
        data["Номер полуприцепа"] = find(r"Номер полуприцепа[:\s]*([^\n\r]+)")
    if not data["ФИО водителя"]:
        data["ФИО водителя"] = find(r"(?:ФИО водителя|Водитель)[:\s]*([^\n\r]+)")

    cust_match = re.search(r"Заказчик:(.*?)(?:Почтовый адрес|$)", text, re.S)
    if cust_match:
        data["Данные заказчика"] = cust_match.group(1).strip()

    inn_match = re.search(r"(ИНН получателя\s*\d+)", text)
    if inn_match:
        data["ИНН получателя"] = inn_match.group(1).strip()

    ogrn_match = re.search(r"(ОГРН\s*\d+)", text)
    if ogrn_match:
        data["ОГРН получателя"] = ogrn_match.group(1).strip()

    return data



def extract_text_from_pdf(path: str) -> str:
    """Extract text from PDF, using OCR for scanned documents."""
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                text += page_text + "\n"
    if text.strip():
        return text

    # Fallback to OCR for scanned PDFs
    images = convert_from_path(path)
    for img in images:
        text += pytesseract.image_to_string(img, lang="rus") + "\n"
    return text


def read_data_from_pdf(path: str):
    text = extract_text_from_pdf(path)
    return parse_data_from_text(text)


def read_data_from_file(path: str):
    if path.lower().endswith(".docx"):
        return read_data_from_docx(path)
    if path.lower().endswith(".pdf"):
        return read_data_from_pdf(path)
    raise ValueError("Unsupported file format")


def replace_placeholders(doc, data):
    """Replace placeholders in doc with values from data."""
    for para in doc.paragraphs:
        for run in para.runs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders(cell, data)


def format_preview(data):
    """Return formatted string for preview widget."""
    lines = [
        data.get("Данные заказчика", ""),
        data.get("ИНН получателя", ""),
        data.get("ОГРН получателя", ""),
        f"Транспортные услуги по договору-заявке {data.get('Номер документа', '')}",
        f"По маршруту {data.get('Адрес загрузки', '')} - {data.get('Адрес разгрузки', '')}",
        f"Автомобиль: {data.get('Марка автомобиля', '')} {data.get('Номер полуприцепа', '')}",
        f"Водитель: {data.get('ФИО водителя', '')}",
        f"Дата погрузки: {data.get('Дата погрузки', '')}",
        f"Дата разгрузки: {data.get('Дата разгрузки', '')}",
        f"Стоимость перевозки: {data.get('Стоимость перевозки', '')}",
    ]
    return "\n".join(lines)


class MainWindow(QtWidgets.QWidget):
    def __init__(self):
        super().__init__()
        self.data = {}
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Word Copywriter")
        layout = QtWidgets.QGridLayout()

        # Source document
        self.source_edit = QtWidgets.QLineEdit()
        self.source_edit.setReadOnly(True)
        source_btn = QtWidgets.QPushButton("Browse")
        source_btn.clicked.connect(self.browse_source)
        layout.addWidget(QtWidgets.QLabel("Source"), 0, 0)
        layout.addWidget(self.source_edit, 1, 0)
        layout.addWidget(source_btn, 2, 0)
        self.preview_edit = QtWidgets.QTextEdit()
        self.preview_edit.setReadOnly(True)
        layout.addWidget(self.preview_edit, 3, 0)

        # Template document
        self.template_edit = QtWidgets.QLineEdit()
        self.template_edit.setReadOnly(True)
        template_btn = QtWidgets.QPushButton("Browse")
        template_btn.clicked.connect(self.browse_template)
        layout.addWidget(QtWidgets.QLabel("Template"), 0, 1)
        layout.addWidget(self.template_edit, 1, 1)
        layout.addWidget(template_btn, 2, 1)

        # Save button
        self.save_btn = QtWidgets.QPushButton("Save")
        self.save_btn.clicked.connect(self.save_document)
        self.save_btn.setEnabled(False)
        layout.addWidget(self.save_btn, 3, 1)

        self.setLayout(layout)

    def browse_source(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Select source", filter="Documents (*.docx *.pdf)"
        )
        if path:
            self.source_edit.setText(path)
            self.data = read_data_from_file(path)
            self.preview_edit.setPlainText(format_preview(self.data))
        self.update_save_button_state()

    def browse_template(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select template", filter="Word Documents (*.docx)")
        if path:
            self.template_edit.setText(path)
        self.update_save_button_state()

    def save_document(self):
        source_path = self.source_edit.text()
        template_path = self.template_edit.text()
        if not (source_path and template_path):
            QMessageBox.warning(self, "Warning", "Please select source and template files")
            return
        output_path, _ = QFileDialog.getSaveFileName(self, "Save document", filter="Word Documents (*.docx)")
        if not output_path:
            return
        if not output_path.lower().endswith('.docx'):
            output_path += '.docx'
        data = getattr(self, 'data', None)
        if not data:
            data = read_data_from_file(source_path)
        doc = Document(template_path)
        replace_placeholders(doc, data)
        try:
            doc.save(output_path)
            QMessageBox.information(self, "Success", f"Document saved to {output_path}")
        except PermissionError:
            QMessageBox.critical(
                self,
                "Error",
                "Cannot save file. It may be open in another program."
            )
            return
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save file: {e}")
            return

    def update_save_button_state(self):
        if self.source_edit.text() and self.template_edit.text():
            self.save_btn.setEnabled(True)
        else:
            self.save_btn.setEnabled(False)


def main():
    app = QtWidgets.QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
